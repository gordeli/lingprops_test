#!/usr/bin/env python3
"""Generate the Web Appendix document for the concreteness library.

Produces a .docx file in a format suitable for a Marketing Science
Web Appendix, documenting the concreteness measure, its normalization,
and the effect of word repetitions.
"""
import numpy as np
from scipy.special import comb
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_text(cell, text, bold=False, size=9):
    """Helper to format a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Times New Roman"

def add_table_row(table, values, bold=False, size=9):
    row = table.add_row()
    for i, val in enumerate(values):
        set_cell_text(row.cells[i], str(val), bold=bold, size=size)
    return row

def main():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # =====================================================================
    # TITLE
    # =====================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Web Appendix: Concreteness Measure — Definition, Normalization, and Properties")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    # =====================================================================
    # WA.1  DEFINITION OF THE CONCRETENESS MEASURE
    # =====================================================================
    h1 = doc.add_heading("WA.1  Definition of the Concreteness Measure", level=1)
    for run in h1.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "We define a concreteness measure for text based on the semantic "
        "specificity of its constituent words, as captured by the WordNet "
        "hypernym hierarchy (Miller 1995; Fellbaum 1998). The measure operates "
        "on content words — nouns, verbs, adjectives, and adverbs — each of "
        "which is mapped to a noun lemma through morphological derivation and "
        "WordNet relations. The depth of each noun lemma in the hypernym tree "
        "serves as a proxy for its concreteness: deeper words (e.g., "
        "\"dachshund\" at depth 13) are more specific and concrete than "
        "shallower words (e.g., \"entity\" at depth 0)."
    )

    doc.add_paragraph(
        "For a text containing K content words (with non-zero concreteness), "
        "let d_i denote the hypernym depth of word i and f_i its frequency "
        "(number of occurrences). When repetitions are counted (the default), "
        "the raw concreteness score is:"
    )

    # Formula
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run(
        "C(text) = \u2211\u1d62 log \u0043(d\u1d62 + f\u1d62, f\u1d62)"
    )
    run.font.size = Pt(11)
    run.italic = True
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "where C(n, k) = n! / [k!(n\u2212k)!] is the binomial coefficient, "
        "and the summation runs over all unique word types i = 1, ..., K "
        "that have a non-zero hypernym depth (d\u1d62 > 0). "
        "Note that the depth value used in the formula is incremented by 1 "
        "relative to the raw hypernym count, so the effective argument is "
        "(d\u1d62 + 1 + f\u1d62 \u2212 1) = (d\u1d62 + f\u1d62)."
    )

    doc.add_paragraph(
        "When a word appears exactly once (f\u1d62 = 1), its contribution "
        "simplifies to log(d\u1d62 + 1). When repetitions are disabled, "
        "each unique noun lemma is counted once regardless of how many "
        "different wordforms map to it."
    )

    # =====================================================================
    # WA.2  NORMALIZATION
    # =====================================================================
    h2 = doc.add_heading("WA.2  Normalization", level=1)
    for run in h2.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "The raw concreteness score C(text) is additive in the number of "
        "contributing words: longer texts mechanically receive higher scores. "
        "To obtain a length-independent measure, we normalize by the "
        "normalization count N, defined as the total number of word tokens "
        "that contribute a non-zero amount to the concreteness sum. "
        "Crucially, words whose hypernym depth is zero (or that cannot be "
        "mapped to a WordNet noun) are excluded from both the numerator and "
        "the denominator:"
    )

    formula2 = doc.add_paragraph()
    formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula2.add_run(
        "C\u0304(text) = C(text) / N"
    )
    run.font.size = Pt(11)
    run.italic = True
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "where N = \u2211\u1d62 f\u1d62 counts only the word tokens (including "
        "repetitions) whose concreteness contribution is strictly positive. "
        "This ensures that function words, stopwords, and content words not "
        "found in WordNet do not dilute the normalized measure."
    )

    doc.add_paragraph(
        "The library provides this normalization count (N) alongside the raw "
        "score for each part-of-speech category (nouns, verbs, adjectives, "
        "adverbs, cardinal numbers) as well as the aggregate total. "
        "Additionally, the library reports the total word count and content-word "
        "counts per POS category, enabling researchers to compute alternative "
        "normalizations if desired."
    )

    # =====================================================================
    # WA.3  EFFECT OF WORD REPETITIONS
    # =====================================================================
    h3 = doc.add_heading("WA.3  Effect of Word Repetitions on the Concreteness Measure", level=1)
    for run in h3.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "The binomial-coefficient formulation creates a specific interaction "
        "between hypernym depth and word frequency. We analyze how repetitions "
        "of the same word affect both the raw and normalized scores."
    )

    # --- WA.3.1 ---
    h31 = doc.add_heading("WA.3.1  The Logarithmic Contribution", level=2)
    for run in h31.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "For a single word with depth d appearing f times, its contribution is:"
    )

    formula3 = doc.add_paragraph()
    formula3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula3.add_run(
        "c(d, f) = log C(d + f, f) = log [(d + f)! / (f! \u00b7 d!)]"
    )
    run.font.size = Pt(11)
    run.italic = True
    run.font.name = "Times New Roman"

    doc.add_paragraph(
        "Using Stirling's approximation for large arguments, this behaves as:"
    )

    cases = doc.add_paragraph()
    cases.paragraph_format.left_indent = Inches(0.5)
    cases.add_run("(i)  f = 1: ").bold = True
    cases.add_run("c(d, 1) = log(d + 1). The contribution is simply the "
                  "logarithm of (depth + 1), growing slowly with depth.\n")
    run2 = cases.add_run("(ii) f \u226b d: ")
    run2.bold = True
    cases.add_run("c(d, f) \u2248 d \u00b7 log(f). The contribution grows "
                  "logarithmically with frequency, scaled by the depth. "
                  "Repeating a concrete word many times yields diminishing "
                  "marginal returns.\n")
    run3 = cases.add_run("(iii) d \u226b f: ")
    run3.bold = True
    cases.add_run("c(d, f) \u2248 f \u00b7 log(d). The contribution grows "
                  "linearly in frequency but only logarithmically in depth.")

    # --- Table 1: Contribution as function of frequency ---
    doc.add_paragraph(
        "Table WA.1 illustrates these properties for a word with hypernym "
        "depth d = 5."
    )

    table1 = doc.add_table(rows=1, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = "Table Grid"

    # Header
    headers = ["Frequency (f)", "C(d+f, f)", "log C(d+f, f)", "Per-word: log C / f"]
    for i, h in enumerate(headers):
        set_cell_text(table1.rows[0].cells[i], h, bold=True, size=9)

    depth = 5
    for freq in [1, 2, 3, 5, 10, 20, 50]:
        c_val = comb(depth + freq, freq)
        lc = np.log(c_val)
        pw = lc / freq
        add_table_row(table1, [
            str(freq),
            f"{c_val:.0f}",
            f"{lc:.4f}",
            f"{pw:.4f}",
        ])

    caption1 = doc.add_paragraph()
    caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption1.add_run("Table WA.1: Concreteness contribution for a word with depth d = 5.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"

    # --- WA.3.2 ---
    h32 = doc.add_heading("WA.3.2  Impact on the Normalized Concreteness", level=2)
    for run in h32.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "When a word is repeated, both the raw score (numerator) and the "
        "normalization count (denominator) increase. However, the raw score "
        "grows sub-linearly (as log C(d + f, f)), while the normalization "
        "count grows linearly (as f). Therefore, the per-word normalized "
        "contribution c(d, f)/f is a strictly decreasing function of f. "
        "This has an important implication:"
    )

    insight = doc.add_paragraph()
    insight.paragraph_format.left_indent = Inches(0.5)
    run = insight.add_run(
        "Repeating the same word lowers the normalized concreteness score. "
        "A text with diverse, concrete vocabulary achieves a higher "
        "normalized concreteness than a text that repeats the same concrete "
        "word multiple times."
    )
    run.italic = True
    run.font.name = "Times New Roman"

    # --- WA.3.3 Limiting Cases ---
    h33 = doc.add_heading("WA.3.3  Limiting Cases", level=2)
    for run in h33.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "We illustrate the behavior of the normalized concreteness measure "
        "with three limiting cases."
    )

    lc1 = doc.add_paragraph()
    lc1.add_run("Case 1: All unique words. ").bold = True
    lc1.add_run(
        "Consider a text of K distinct content words, each appearing once "
        "(f\u1d62 = 1 for all i). The normalized concreteness is "
        "\u2211 log(d\u1d62 + 1) / K, which is simply the average "
        "log-depth across the vocabulary. This represents the maximum "
        "normalized score achievable for a given set of word types."
    )

    lc2 = doc.add_paragraph()
    lc2.add_run("Case 2: Single word repeated N times. ").bold = True
    lc2.add_run(
        "If a single word with depth d is repeated N times, the normalized "
        "score is log C(d + N, N) / N. As N \u2192 \u221e, this approaches "
        "d \u00b7 log(N) / N \u2192 0. Thus, extreme repetition drives the "
        "normalized concreteness toward zero, regardless of how concrete "
        "the repeated word is."
    )

    lc3 = doc.add_paragraph()
    lc3.add_run("Case 3: Mixed text with some repeated words. ").bold = True
    lc3.add_run(
        "In the general case, words with high frequency contribute less "
        "per-word than words with low frequency. This means the normalized "
        "measure naturally penalizes redundancy and rewards lexical diversity, "
        "which is a desirable property for measuring the richness of concrete "
        "language in text."
    )

    # --- Table 2: Practical examples ---
    doc.add_paragraph(
        "Table WA.2 demonstrates these effects using the library's output "
        "for simple test phrases."
    )

    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = "Table Grid"

    headers2 = ["Text", "Raw Score", "Norm. Count (N)", "Normalized Score"]
    for i, h in enumerate(headers2):
        set_cell_text(table2.rows[0].cells[i], h, bold=True, size=9)

    # Compute examples using the library
    from lingprops import compute_concreteness
    examples = [
        "cat",
        "cat cat cat",
        "cat dog mouse",
        "cat dog",
        "cat cat",
    ]
    for ex in examples:
        r = compute_concreteness(ex)
        t = r["total"]
        add_table_row(table2, [
            f'"{ex}"',
            f'{t["score"]:.4f}',
            str(t["count"]),
            f'{t["normalized_score"]:.4f}',
        ])

    caption2 = doc.add_paragraph()
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption2.add_run(
        'Table WA.2: Concreteness scores for example phrases. '
        'Note that "cat dog" (diverse) achieves a higher normalized score '
        'than "cat cat" (repeated).'
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"

    # =====================================================================
    # WA.4  OUTPUT SPECIFICATION
    # =====================================================================
    h4 = doc.add_heading("WA.4  Library Output Specification", level=1)
    for run in h4.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "The compute_concreteness() function returns a dictionary with the "
        "following structure for each POS category and the aggregate total:"
    )

    spec = doc.add_paragraph()
    spec.paragraph_format.left_indent = Inches(0.5)
    spec.add_run("score").bold = True
    spec.add_run(
        " — The raw (unnormalized) concreteness sum: "
        "\u2211 log C(d\u1d62 + f\u1d62, f\u1d62).\n"
    )
    run_c = spec.add_run("count")
    run_c.bold = True
    spec.add_run(
        " — The normalization count N: total number of word tokens with "
        "non-zero concreteness contribution.\n"
    )
    run_ns = spec.add_run("normalized_score")
    run_ns.bold = True
    spec.add_run(
        " — The normalized concreteness: score / count "
        "(0 when count = 0).\n"
    )
    run_wc = spec.add_run("word_count")
    run_wc.bold = True
    spec.add_run(
        " (total only) — Total number of word tokens in the text.\n"
    )
    run_cwc = spec.add_run("content_word_counts")
    run_cwc.bold = True
    spec.add_run(
        " (total only) — Dictionary of word counts per content POS category "
        "(NN, VB, JJ, RB, CD), counting all words of that POS regardless of "
        "whether they contribute to concreteness.\n"
    )

    doc.add_paragraph(
        "Additionally, the count_words() function provides standalone access "
        "to the word count and per-POS content word counts without computing "
        "concreteness scores."
    )

    # =====================================================================
    # WA.5  MATHEMATICAL PROPERTIES
    # =====================================================================
    h5 = doc.add_heading("WA.5  Mathematical Properties", level=1)
    for run in h5.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "The concreteness measure has the following desirable mathematical "
        "properties:"
    )

    props = [
        ("Non-negativity", "C(text) \u2265 0, with equality when the text "
         "contains no content words mappable to WordNet nouns."),
        ("Additivity", "The raw score is additive across word types: "
         "C(text) = \u2211\u1d62 c(d\u1d62, f\u1d62). This allows decomposition "
         "by POS category."),
        ("Monotonicity in depth", "For fixed frequency, c(d, f) is strictly "
         "increasing in d. More specific (deeper) words contribute more."),
        ("Sub-linearity in frequency", "For fixed depth, c(d, f) grows "
         "sub-linearly in f (as \u0398(d log f) for large f). Repeating a word "
         "yields diminishing marginal returns."),
        ("Diversity preference", "The normalized score c(d, f)/f is strictly "
         "decreasing in f, so diverse vocabulary is rewarded over repetition."),
    ]

    for name, desc in props:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(f"{name}. ")
        run.bold = True
        p.add_run(desc)

    # =====================================================================
    # REFERENCES
    # =====================================================================
    h_ref = doc.add_heading("References", level=1)
    for run in h_ref.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    refs = [
        "Fellbaum, C. (1998). WordNet: An Electronic Lexical Database. "
        "MIT Press, Cambridge, MA.",
        "Miller, G. A. (1995). WordNet: A Lexical Database for English. "
        "Communications of the ACM, 38(11), 39\u201341.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

    # --- Save ---
    out_path = "Web_Appendix_Concreteness.docx"
    doc.save(out_path)
    print(f"Document saved to: {out_path}")


if __name__ == "__main__":
    main()
