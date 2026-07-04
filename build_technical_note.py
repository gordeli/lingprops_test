#!/usr/bin/env python3
"""Build .docx and .pdf versions of the concreteness technical note.

Single source of truth for both outputs (python-docx + reportlab), so the
Word and PDF versions never drift. Content mirrors
`Concreteness_RepNorep_Normalization_WSD.md`.

Outputs (repo root):
    Concreteness_Technical_Note.docx
    Concreteness_Technical_Note.pdf
"""
import re
import os

OUT_DOCX = r"C:\GitHub\lingprops_repackaged\Concreteness_Technical_Note.docx"
OUT_PDF = r"C:\GitHub\lingprops_repackaged\Concreteness_Technical_Note.pdf"

# ---------------------------------------------------------------------------
# CONTENT MODEL
# Each block is a tuple; renderers below interpret them for docx and pdf.
#   ("title", text)
#   ("subtitle", text)
#   ("h1", text)
#   ("h2", text)
#   ("p", text)            inline markup: `code`, **bold**
#   ("formula", text)      centered, italic
#   ("code", text)         monospace block
#   ("table", headers, rows, colspec)  colspec = list of relative widths or None
# ---------------------------------------------------------------------------

BLOCKS = [
    ("title", "Concreteness Scoring: Repetitions, Normalization, "
              "Text-Level Aggregation, and Word-Sense Disambiguation"),
    ("subtitle", "Technical note for coauthors"),

    ("p", "**Scope.** This note documents exactly how the `lingprops` "
          "concreteness scorer treats (1) repeated words, (2) the normalization "
          "denominator, (3) sentence-level vs. whole-text accumulation, and "
          "(4) word senses under the different WSD strategies. Every numerical "
          "value below was produced by running the library; each is reproduced "
          "by a closed-form hand check that matches to machine precision."),
    ("p", "Code references point to `src/lingprops/concreteness.py` and "
          "`src/lingprops/_concreteness_legacy.py`."),

    ("h1", "1. The unit of accounting: the wordform"),
    ("p", "Tokenization happens in `legacy.wordformtion(text)` "
          "(`_concreteness_legacy.py:397`). It runs `sent_tokenize` -> "
          "per-sentence `word_tokenize` -> per-sentence `pos_tag`, but it "
          "accumulates every token into a single dictionary for the whole text:"),
    ("code", "wordforms[(word.lower(), tag)] = wordforms.get((word.lower(), tag), 0) + 1   # line 553"),
    ("p", "Consequently the atomic unit everywhere downstream is a **wordform**: "
          "the pair (lowercased surface form, POS tag), carrying a frequency "
          "**f** counted over the entire text. Sentence boundaries are discarded "
          "after this step. This single design choice determines the answers in "
          "Sections 3 and 4."),

    ("h1", "2. The scoring formula"),
    ("p", "Each POS category (NN, VB, JJ, RB, CD) is scored as an independent "
          "partition. For a wordform of depth **d** (number of transitive "
          "WordNet hypernyms of its chosen synset) and frequency **f**, the "
          "contribution to the partition score is"),
    ("formula", "Delta  =  log C(d + f, f)"),
    ("p", "the logarithm of a multiset coefficient (the number of size-f "
          "multisets drawn from d+1 slots). In code: "
          "`np.log(comb(depth + 1 + frequency - 1, frequency))` "
          "(`concreteness.py:170`, `:183`), which equals log C(d+f, f)."),
    ("p", "A wordform is **valid** (contributes to both numerator and "
          "denominator) only if it maps to a WordNet noun-lemma, is not "
          "excluded, and has d > 0 (the sole exception being the literal lemma "
          "\"entity\"). Words carrying zero concreteness therefore do not "
          "inflate the denominator (`concreteness.py:168`, `:181`)."),
    ("h2", "Two variants"),
    ("table",
     ["Variant", "Frequency used", "Per-wordform term", "Denominator increment"],
     [["With repetitions (rep)", "actual f", "log C(d+f, f)", "+ f"],
      ["Without repetitions (norep)", "forced f = 1", "log C(d+1, 1) = log(d+1)", "+ 1"]],
     [1.4, 1.1, 1.6, 1.1]),
    ("p", "**rep** - `_compute_pos_score` (`concreteness.py:189`): iterates "
          "wordforms, uses the real whole-text count f, adds f to the "
          "denominator."),
    ("p", "**norep** - `_compute_pos_score_norep` (`concreteness.py:214`): "
          "deduplicates by lemma within the POS partition (WordNetLemmatizer "
          "output, before nounification; `concreteness.py:235`), scores each "
          "unique lemma once with f = 1, adds 1 to the denominator."),
    ("h2", "Normalization (the denominator)"),
    ("p", "**rep denominator** = sum of f over valid wordforms in the partition "
          "(`concreteness.py:209`)."),
    ("p", "**norep denominator** = number of valid unique lemmas in the "
          "partition (`concreteness.py:247`)."),
    ("p", "Final normalized values: normalized_score = score / count and "
          "normalized_score_norep = score_norep / count_norep "
          "(`concreteness.py:408`, `:411`)."),

    ("h1", "3. Sentence-level vs. whole-text: repeated words are NOT double-counted"),
    ("p", "Because of the whole-text aggregation in Section 1, scoring is "
          "effectively whole-text, not per-sentence. A concern that would apply "
          "to a per-sentence implementation - the same word reappearing in a "
          "second sentence contributing a second time and inflating the "
          "denominator - does not occur here:"),
    ("p", "- A noun in three sentences (all NN) becomes one entry "
          "('cat','NN'): 3."),
    ("p", "- **rep**: scored once as log C(d+3, 3); denominator += 3."),
    ("p", "- **norep**: deduped to one lemma -> log(d+1); denominator += 1."),
    ("p", "This is demonstrated by the identity of Case 2 and Case 3 in "
          "Section 5."),

    ("h1", "4. Word-sense disambiguation: senses are NOT treated separately"),
    ("p", "WSD (strategies `first`, `lesk`, `neural`; see "
          "`src/lingprops/wsd.py`) only selects which single depth a wordform "
          "receives. It never splits a wordform into per-sense entries and never "
          "changes the counts:"),
    ("p", "- The picker is called once per wordform (`concreteness.py:141`) and "
          "returns one synset -> one depth."),
    ("p", "- Its context is the whole text "
          "(`context_tokens = nltk.word_tokenize(text)`, `concreteness.py:365`; "
          "the neural picker even caches one context embedding per document)."),
    ("p", "So a word like `bank` occurring twice with two different senses is "
          "still one wordform ('bank','NN'). Its rep denominator counts its "
          "frequency; its norep denominator counts it once. Changing the WSD "
          "strategy can change the score (a different single depth is chosen) "
          "but cannot change either denominator."),
    ("p", "**Implication.** True per-occurrence, per-sentence sense "
          "disambiguation is not possible in the current architecture, because "
          "`wordformtion` discards token positions and sentence membership "
          "before scoring. All occurrences of one (word, tag) are "
          "indistinguishable and share a single sense/depth."),

    ("h1", "5. Worked numerical demo"),
    ("p", "Noun (NN) partition only, `ner=False`. Depths from the real WordNet "
          "hierarchy under the default strategy:"),
    ("table",
     ["word", "lemma", "depth d"],
     [["cat", "cat", "13"], ["dog", "dog", "14"], ["bank", "bank", "5"]],
     [1, 1, 1]),
    ("p", "Per-wordform term: log C(d+f, f); norep term collapses to log(d+1)."),

    ("h2", "5.1 Scores and denominators"),
    ("table",
     ["Case", "Text", "REP\nscore", "REP\ndenom", "REP\nnorm",
      "NOREP\nscore", "NOREP\ndenom", "NOREP\nnorm"],
     [["1 - cat x1", "The cat slept.", "2.6391", "1", "2.6391",
       "2.6391", "1", "2.6391"],
      ["2 - cat x3, one sentence", "The cat saw the cat near the cat.",
       "6.3279", "3", "2.1093", "2.6391", "1", "2.6391"],
      ["3 - cat x3, three sentences",
       "The cat slept. A cat ran. My cat purred.",
       "6.3279", "3", "2.1093", "2.6391", "1", "2.6391"],
      ["4 - cat + dog, distinct", "The cat and the dog played.",
       "5.3471", "2", "2.6736", "5.3471", "2", "2.6736"],
      ["5 - cat x3 + dog x1",
       "The cat chased a cat while another cat watched the dog.",
       "9.0360", "4", "2.2590", "5.3471", "2", "2.6736"]],
     [1.5, 2.6, 0.7, 0.7, 0.7, 0.8, 0.8, 0.8]),
    ("p", "**Hand checks (all match the library output):**"),
    ("p", "- Case 1: log C(14,1) = log 14 = 2.6391. No repetition => rep = norep."),
    ("p", "- Case 2: rep log C(16,3) = 6.3279, denom 3. norep log 14 = 2.6391, "
          "denom 1."),
    ("p", "- **Case 3 is bit-identical to Case 2** - the decisive illustration "
          "that aggregation is whole-text: three cats across three different "
          "sentences leave the score, both denominators, and both normalized "
          "values unchanged."),
    ("p", "- Case 4: log 14 + log 15 = 5.3471, denom 2. No repetition => rep = norep."),
    ("p", "- Case 5: rep log C(16,3) + log C(15,1) = 9.0360, denom 3+1 = 4. "
          "norep log 14 + log 15 = 5.3471, denom 1+1 = 2."),

    ("h2", "5.2 How repetition moves the REP term (fixed depth d = 13)"),
    ("table",
     ["f", "REP term  log C(d+f,f)", "REP normalized  (term/f)", "NOREP term"],
     [["1", "2.6391", "2.6391", "2.6391"],
      ["2", "4.6540", "2.3270", "2.6391"],
      ["3", "6.3279", "2.1093", "2.6391"],
      ["5", "9.0558", "1.8112", "2.6391"],
      ["10", "13.9501", "1.3950", "2.6391"],
      ["50", "29.9794", "0.5996", "2.6391"]],
     [0.6, 1.6, 1.6, 1.1]),
    ("p", "The multiset coefficient grows sublinearly in f, while the "
          "denominator grows linearly (+f). Hence the rep-normalized "
          "contribution of a word decreases the more it is repeated - "
          "repetition dilutes per-word concreteness. The norep term is flat by "
          "construction (log(d+1)), fully immune to repetition."),
    ("p", "**Limiting cases.** f = 1: rep = norep exactly (Cases 1, 4). "
          "f -> infinity: log C(d+f, f) ~ d * log f grows without bound but only "
          "logarithmically, so rep-normalized -> 0; norep is unchanged at "
          "log(d+1)."),

    ("h2", "5.3 WSD does not touch the denominators"),
    ("p", "Text: \"I sat by the river bank. Then I went to the bank to deposit "
          "cash.\" (nouns: river, bank x2, cash)"),
    ("table",
     ["Strategy", "REP score", "REP denom", "NOREP score", "NOREP denom"],
     [["first", "6.9157", "4", "5.6630", "3"],
      ["lesk", "6.9157", "4", "5.6630", "3"]],
     [1, 1, 1, 1, 1]),
    ("p", "`bank` appears twice with two plausible senses but is one wordform. "
          "rep denom = 1(river) + 2(bank) + 1(cash) = 4; norep denom = 3 unique "
          "lemmas - identical across strategies. Here first and lesk also "
          "happened to select the same depth, so even the scores coincide; in "
          "general only the score can differ, never the counts."),

    ("h1", "6. Summary"),
    ("p", "1. **rep vs. norep.** rep scores each wordform once as log C(d+f, f) "
          "with the true whole-text frequency f and adds f to the denominator; "
          "norep deduplicates to unique lemmas per POS, scores each as "
          "log(d+1), and adds 1."),
    ("p", "2. **Normalization.** The denominator counts only words with "
          "non-zero concreteness - sum of f (rep) or the number of unique valid "
          "lemmas (norep)."),
    ("p", "3. **Text-level, not sentence-level.** Counts are aggregated over "
          "the whole text, so a word repeated across sentences is scored once "
          "(as a multiset of size f) and never double-counts the numerator or "
          "the denominator."),
    ("p", "4. **WSD.** Disambiguation acts per wordform using whole-text "
          "context and only sets a single depth. Different senses of one surface "
          "form are the same unit; WSD changes the score but never the rep/norep "
          "denominators."),
    ("p", "Reproduce with the accompanying script `demo_repnorep.py`."),
]


# ---------------------------------------------------------------------------
# Inline markup tokenizer: splits text into (style, text) segments.
# Supports `code` and **bold** (non-nested).
# ---------------------------------------------------------------------------
_TOKEN = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def tokenize_inline(text):
    segments = []
    for part in _TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            segments.append(("bold", part[2:-2]))
        elif part.startswith("`") and part.endswith("`"):
            segments.append(("code", part[1:-1]))
        else:
            segments.append(("normal", part))
    return segments


# ===========================================================================
# DOCX RENDERER
# ===========================================================================
def build_docx():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, m, Cm(2.2))

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    def add_runs(p, text, base_size=10.5):
        for style, seg in tokenize_inline(text):
            run = p.add_run(seg)
            run.font.size = Pt(base_size)
            if style == "bold":
                run.bold = True
                run.font.name = "Calibri"
            elif style == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(base_size - 1)
                run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
            else:
                run.font.name = "Calibri"

    def set_cell(cell, text, bold=False, size=8.5, header=False):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold or header

    for block in BLOCKS:
        kind = block[0]
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(block[1]); r.bold = True
            r.font.size = Pt(15); r.font.name = "Calibri"
        elif kind == "subtitle":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(block[1]); r.italic = True
            r.font.size = Pt(11); r.font.name = "Calibri"
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            doc.add_paragraph()
        elif kind == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r = p.add_run(block[1]); r.bold = True
            r.font.size = Pt(13); r.font.name = "Calibri"
            r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x5C)
        elif kind == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(block[1]); r.bold = True
            r.font.size = Pt(11.5); r.font.name = "Calibri"
            r.font.color.rgb = RGBColor(0x2E, 0x5A, 0x88)
        elif kind == "p":
            p = doc.add_paragraph()
            add_runs(p, block[1])
        elif kind == "formula":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(block[1]); r.italic = True
            r.font.size = Pt(12); r.font.name = "Cambria Math"
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(block[1])
            r.font.name = "Consolas"; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif kind == "table":
            headers, rows, _colspec = block[1], block[2], block[3]
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Light Grid Accent 1"
            for i, h in enumerate(headers):
                set_cell(t.rows[0].cells[i], h.replace("\n", " "),
                         header=True, size=8.5)
            for row in rows:
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    set_cell(cells[i], str(val), size=8.5)
            doc.add_paragraph()

    doc.save(OUT_DOCX)
    print("wrote", OUT_DOCX)


# ===========================================================================
# PDF RENDERER (reportlab)
# ===========================================================================
def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, Preformatted)
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["Normal"], fontName="Helvetica",
                          fontSize=10, leading=14, spaceAfter=5, alignment=TA_LEFT)
    title = ParagraphStyle("title", parent=styles["Normal"], fontName="Helvetica-Bold",
                           fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=4)
    subtitle = ParagraphStyle("subtitle", parent=styles["Normal"],
                              fontName="Helvetica-Oblique", fontSize=11,
                              leading=14, alignment=TA_CENTER, spaceAfter=14,
                              textColor=colors.HexColor("#555555"))
    h1 = ParagraphStyle("h1", parent=styles["Normal"], fontName="Helvetica-Bold",
                        fontSize=13, leading=16, spaceBefore=12, spaceAfter=5,
                        textColor=colors.HexColor("#1F3B5C"))
    h2 = ParagraphStyle("h2", parent=styles["Normal"], fontName="Helvetica-Bold",
                        fontSize=11.5, leading=14, spaceBefore=8, spaceAfter=4,
                        textColor=colors.HexColor("#2E5A88"))
    formula = ParagraphStyle("formula", parent=styles["Normal"],
                             fontName="Helvetica-Oblique", fontSize=12,
                             leading=16, alignment=TA_CENTER, spaceBefore=4,
                             spaceAfter=8)
    codestyle = ParagraphStyle("code", parent=styles["Normal"],
                               fontName="Courier", fontSize=8.5, leading=11,
                               leftIndent=10, textColor=colors.HexColor("#333333"),
                               backColor=colors.HexColor("#F3F3F3"),
                               spaceBefore=2, spaceAfter=6)
    cell_style = ParagraphStyle("cell", parent=styles["Normal"],
                                fontName="Helvetica", fontSize=8, leading=10)
    cell_head = ParagraphStyle("cellh", parent=cell_style,
                               fontName="Helvetica-Bold", textColor=colors.white)

    def inline_to_rl(text):
        out = []
        for style, seg in tokenize_inline(text):
            seg = escape(seg)
            if style == "bold":
                out.append("<b>%s</b>" % seg)
            elif style == "code":
                out.append('<font face="Courier" color="#8B0000">%s</font>' % seg)
            else:
                out.append(seg)
        return "".join(out)

    story = []
    avail_w = A4[0] - 2 * 1.8 * cm

    for block in BLOCKS:
        kind = block[0]
        if kind == "title":
            story.append(Paragraph(escape(block[1]), title))
        elif kind == "subtitle":
            story.append(Paragraph(escape(block[1]), subtitle))
        elif kind == "h1":
            story.append(Paragraph(escape(block[1]), h1))
        elif kind == "h2":
            story.append(Paragraph(escape(block[1]), h2))
        elif kind == "p":
            story.append(Paragraph(inline_to_rl(block[1]), body))
        elif kind == "formula":
            story.append(Paragraph(escape(block[1]), formula))
        elif kind == "code":
            story.append(Preformatted(block[1], codestyle))
        elif kind == "table":
            headers, rows, colspec = block[1], block[2], block[3]
            if colspec is None:
                colspec = [1] * len(headers)
            total = sum(colspec)
            widths = [avail_w * w / total for w in colspec]
            data = [[Paragraph(escape(h.replace("\n", " ")), cell_head)
                     for h in headers]]
            for row in rows:
                data.append([Paragraph(escape(str(v)), cell_style) for v in row])
            tbl = Table(data, colWidths=widths, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2E5A88")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.HexColor("#F2F6FA")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 8))

    doc = SimpleDocTemplate(OUT_PDF, pagesize=A4,
                            topMargin=1.8 * cm, bottomMargin=1.8 * cm,
                            leftMargin=1.8 * cm, rightMargin=1.8 * cm,
                            title="Concreteness Technical Note")
    doc.build(story)
    print("wrote", OUT_PDF)


if __name__ == "__main__":
    build_docx()
    build_pdf()
